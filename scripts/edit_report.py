"""One-shot editor: take the CU MCA template and produce the EduPulse first-pass report.

Pass 1 (this run):
  - Fix margins to 1" all around (guideline)
  - Force Times New Roman + 12 pt + 1.5 line spacing on default style
  - Fix Chapter 4 heading style (Normal → Heading 1)
  - Fix Chapter 5 typo (Implémentation → Implementation)
  - Fill Declaration placeholders with personal details
  - Draft Acknowledgement
  - Draft Abstract (using real model metrics)

Output:
  EduPulse_Report.docx  (next to the template — template is left untouched)
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Project Report Template.docx"
DST = ROOT / "EduPulse_Report.docx"

# ----------------------------------------------------------------- inputs

STUDENT_NAME = "Deepak Kumar Joshi"
ENROLLMENT = "O24MCA112387"
PROJECT_TITLE = (
    "EduPulse: A Business Intelligence Dashboard for "
    "Online Learning Engagement Analytics"
)
ACADEMIC_YEAR = "2025–2026"
SUPERVISOR = "Vikas Kumar, Senior Mentor"
PLACE = "Hyderabad, Telangana"
PROGRAM = "Master of Computer Applications (Online)"
INSTITUTION = "Chandigarh University (CU Online), Chandigarh"

# ---------------------------------------------------------- helpers

def set_font_everywhere(rPr_element, font_name: str) -> None:
    """Set ascii / hAnsi / eastAsia / cs font names on an rFonts element."""
    rFonts = rPr_element.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr_element.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


def force_style_font(style, font_name: str = "Times New Roman") -> None:
    """Force a style to use Times New Roman everywhere (incl. East Asian / complex)."""
    style.font.name = font_name
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        from docx.oxml import OxmlElement

        rPr = OxmlElement("w:rPr")
        style.element.append(rPr)
    set_font_everywhere(rPr, font_name)


def replace_paragraph_text(p, new_text: str) -> None:
    """Replace a paragraph's text while keeping the first run's formatting."""
    if not p.runs:
        p.add_run(new_text)
        return
    first = p.runs[0]
    first.text = new_text
    for r in p.runs[1:]:
        r.text = ""


def append_paragraph(doc, text: str, style_name: str = "Normal"):
    p = doc.add_paragraph(text)
    p.style = doc.styles[style_name]
    return p


# ----------------------------------------------------------------- run

def main() -> None:
    doc = Document(str(SRC))

    # ---------------- 1. margins to 1" -----------------
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ---------------- 2. base font + spacing -----------------
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3",
                       "Title", "Subtitle", "List Paragraph", "Normal (Web)"):
        if style_name in doc.styles:
            force_style_font(doc.styles[style_name])

    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Heading 1: 14-16pt, bold — already 14pt in template, keep but ensure TNR is set
    # Heading 2: 13pt — bump to 14 for clarity
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 2"].font.bold = True

    # ---------------- 3. fix template bugs -----------------
    # Find paragraphs by their text content (indices may shift)
    paragraphs = doc.paragraphs

    # Chapter 4 heading uses Normal style → Heading 1
    for p in paragraphs:
        if p.text.strip().lower().startswith("chapter 4: system design"):
            p.style = doc.styles["Heading 1"]
            break

    # Chapter 5 typo
    for p in paragraphs:
        if "Implémentation" in p.text:
            replace_paragraph_text(p, p.text.replace("Implémentation", "Implementation"))

    # ---------------- 4. fill Declaration -----------------
    declaration_lines = {
        # match by partial unique prefix → replacement
        "I, ____________________, hereby solemnly declare":
            f"I, {STUDENT_NAME} (Enrollment No: {ENROLLMENT}), hereby solemnly declare "
            f"that the project report titled “{PROJECT_TITLE}” submitted in "
            f"partial fulfilment of the requirements for the award of the degree of "
            f"{PROGRAM} at {INSTITUTION} is an original record of my own work.",
        "This project has been carried out by me during the academic year":
            f"This project has been carried out by me during the academic year "
            f"{ACADEMIC_YEAR} under the supervision of {SUPERVISOR}.",
        "The work has not been submitted previously":
            "The work has not been submitted previously to any other university, "
            "institution, or examination body for the award of any degree, diploma, "
            "or other similar title.",
        "All sources of information used in this report":
            "All sources of information used in this report have been duly acknowledged "
            "and referenced in accordance with academic ethics.",
        "The data presented in this report":
            "The data presented in this report is authentic to the best of my knowledge "
            "and has not been fabricated or manipulated.",
        "I understand that if any part of this declaration":
            "I understand that if any part of this declaration is found to be false, "
            "my project may be rejected and disciplinary action may be initiated as per "
            "the rules of the university.",
        "Place: ____________________":
            f"Place: {PLACE}\nDate: ____________________",
        "Student Signature:":
            f"Student Signature: ____________________\n"
            f"Student Name: {STUDENT_NAME}\n"
            f"Enrollment No: {ENROLLMENT}",
    }
    for p in doc.paragraphs:
        for key, val in declaration_lines.items():
            if key in p.text:
                replace_paragraph_text(p, val)
                break

    # ---------------- 5. Acknowledgement -----------------
    ack_text = (
        f"I take this opportunity to express my sincere gratitude to my project guide "
        f"{SUPERVISOR}, for the invaluable mentorship, technical guidance, and constant "
        f"encouragement provided throughout the development of this project. His "
        f"insights on industry-grade analytics workflows and structured project "
        f"execution shaped both the scope and the quality of EduPulse.\n\n"
        f"I extend my heartfelt thanks to the faculty members of the Department of "
        f"Computer Applications at {INSTITUTION}, whose curriculum and timely guidance "
        f"laid the foundation that made this project possible. I am also grateful to "
        f"the Open University, United Kingdom, for releasing the OULAD dataset under "
        f"the CC BY 4.0 licence, which served as the empirical backbone of this work.\n\n"
        f"Finally, I thank my family and friends for their unwavering moral support "
        f"during my MCA journey, and my peers for the constructive discussions that "
        f"helped sharpen many design decisions in this report."
    )
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip() == "Acknowledgement":
            # Replace the next paragraph (the placeholder) with the full ack text
            next_p = doc.paragraphs[i + 1]
            replace_paragraph_text(next_p, ack_text)
            break

    # ---------------- 6. Abstract -----------------
    abstract_text = (
        f"This project, titled “{PROJECT_TITLE}”, develops an end-to-end "
        f"Business Intelligence platform that measures learner engagement on online "
        f"learning environments and identifies students at risk of withdrawing from a "
        f"course before completion. The system addresses a critical gap in online "
        f"learning operations—administrators rarely have access to consolidated, "
        f"interactive engagement analytics, and individual at-risk learners are usually "
        f"detected only after they have already disengaged. EduPulse ingests the "
        f"anonymised Open University Learning Analytics Dataset (32,593 students, 22 "
        f"course presentations, over ten million click interactions), cleans and "
        f"validates it through a reproducible Python pipeline, and loads it into a "
        f"DuckDB star-schema data warehouse comprising three fact tables and four "
        f"conformed dimensions. A Streamlit multi-page dashboard then surfaces "
        f"engagement trends, course popularity, high- and low-engagement periods, "
        f"and category-level comparisons through interactive Plotly visualisations. "
        f"A supervised Gradient Boosting classifier, trained on first-thirty-day "
        f"engagement and demographic features, predicts learner withdrawal with "
        f"80.3% accuracy and 0.810 ROC AUC on a held-out test set of 6,519 "
        f"enrollments. Development followed an iterative, milestone-based methodology "
        f"covering data collection, cleansing, warehouse design, ETL implementation, "
        f"dashboard development, and validation. The complete solution is deployed "
        f"free of cost on Streamlit Community Cloud, providing administrators a "
        f"single web interface for both descriptive analytics and predictive early-"
        f"warning of dropout risk."
    )
    for i, p in enumerate(doc.paragraphs):
        if (p.style.name == "Heading 1" and p.text.strip().lower() == "abstract"
                and i > 25):  # the second Abstract heading, not the executive-summary one
            # Replace the "(150-250 words)" placeholder + "Write content here..." with abstract
            placeholder = doc.paragraphs[i + 1]
            replace_paragraph_text(placeholder, abstract_text)
            # Clear the "Write content here..." line that follows
            try:
                following = doc.paragraphs[i + 2]
                if "Write content here" in following.text:
                    replace_paragraph_text(following, "")
            except IndexError:
                pass
            break

    # ---------------- save -----------------
    doc.save(str(DST))
    print(f"Wrote {DST}")
    print(f"Paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
